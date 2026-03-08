#!/usr/bin/env python3
"""Generate Dr. Bronner's Customer Analysis Report for Priority1 Logistics."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Style Configuration ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'  # Clean sans-serif fallback (Inter not typically available in docx)
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Configure heading styles
for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    elif level == 3:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    else:
        h.font.size = Pt(11)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_thin_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'Arial'
        p.add_run(text).font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = 'Arial'


def add_table_row_shading(row, color):
    """Apply shading to a table row."""
    for cell in row.cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)


def make_table(doc, headers, rows):
    """Create a clean formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    add_table_row_shading(hdr, "2D2D2D")
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        if r_idx % 2 == 1:
            add_table_row_shading(row, "F5F5F5")
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'

    return table


# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CUSTOMER ANALYSIS")
run.font.size = Pt(32)
run.font.name = 'Arial'
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Dr. Bronner's Magic Soaps")
run.font.size = Pt(20)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()

prep = doc.add_paragraph()
prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prep.add_run("Prepared by Priority1 Logistics")
run.font.size = Pt(12)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("March 2026")
run.font.size = Pt(12)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
conf.paragraph_format.space_before = Pt(40)
run = conf.add_run("CONFIDENTIAL \u2014 FOR INTERNAL USE ONLY")
run.font.size = Pt(9)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Full Prospect Intelligence Brief",
    "2. Products and Freight Profile Translation",
    "3. Supply Chain and Network Mapping",
    "4. Industry-Specific Logistics Challenges",
    "5. Recent News and Trigger Events Analysis",
    "6. Competitive Landscape and Broker Opportunity",
    "7. Sales Strategy and Messaging Framework",
    "8. Executive-Level Logistics Summary",
    "Appendix: Internal Deal Brief",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 1: FULL PROSPECT INTELLIGENCE BRIEF
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Full Prospect Intelligence Brief", level=1)
add_thin_line(doc)

doc.add_heading("Company Overview", level=2)
doc.add_paragraph(
    "Dr. Bronner's Magic Soaps is a privately held, family-owned manufacturer of organic and fair trade "
    "personal care products, headquartered in Vista, California. Founded in 1948 by Emanuel Bronner, the "
    "company is now led by CEO David Bronner (grandson of the founder) and COO Michael Bronner. Dr. Bronner's "
    "is the top-selling natural brand of soap in the United States and has grown from approximately $4 million "
    "in annual revenue in 1998 to an estimated $228 million in recent years, representing extraordinary organic "
    "growth driven by brand loyalty, values alignment, and word-of-mouth rather than traditional advertising."
)
doc.add_paragraph(
    "The company employs approximately 300-500 people across its operations, with the bulk of its workforce "
    "located at its Vista, California manufacturing headquarters. Dr. Bronner's operates across five continents "
    "and ships finished products to more than 40 countries worldwide. As of January 2026, Dr. Bronner's is a "
    "Certified Living Wage employer, with a starting hourly wage of $29.11 and a 5-to-1 cap on executive-to-"
    "worker pay ratios."
)

doc.add_heading("Products and Services", level=2)
doc.add_paragraph(
    "Dr. Bronner's product portfolio centers on its iconic 18-in-1 Pure-Castile Liquid Soap, available in "
    "eight scents and multiple sizes (2 oz travel, 8 oz, 16 oz, 32 oz, 32 oz refill carton, and 1 gallon). "
    "Beyond liquid soap, the product line includes:"
)
add_bullet(doc, "Pure-Castile Bar Soaps (5 oz bars, eight scent varieties)")
add_bullet(doc, "All-One Toothpaste (fluoride-free, 1 oz travel and 5 oz sizes)")
add_bullet(doc, "Organic Lip Balms (0.15 oz)")
add_bullet(doc, "Organic Hand Sanitizing Sprays (2 oz)")
add_bullet(doc, "Organic Hair Rinse and Hair Creme")
add_bullet(doc, "Organic Body Lotions and Shaving Soaps")
add_bullet(doc, "Sal Suds Biodegradable Cleaner (household cleaning concentrate)")
add_bullet(doc, "Whole Kernel Coconut Oil (30 oz)")
add_bullet(doc, "Soap Refill Cartons (launched into nearly 10,000 retail stores nationally)")

doc.add_paragraph(
    "All products are certified organic under the USDA National Organic Program, fair trade certified under "
    "the Fair for Life programme, non-GMO, vegan (with one exception), and cruelty-free. The brand is known for "
    "its dense, text-heavy labels featuring philosophical and social messages."
)

doc.add_heading("Industries and End Markets", level=2)
doc.add_paragraph(
    "Dr. Bronner's operates primarily in the personal care and consumer packaged goods (CPG) sector, with "
    "exposure to the following end markets:"
)
add_bullet(doc, "Natural and organic grocery (Whole Foods Market, Sprouts, Natural Grocers, NCG, INFRA, Fresh Thyme)")
add_bullet(doc, "Mass market retail (Target, Walmart, Costco, Kroger, Trader Joe's)")
add_bullet(doc, "Specialty retail (REI, ULTA Beauty)")
add_bullet(doc, "Drug and pharmacy chains (Walgreens)")
add_bullet(doc, "E-commerce (drbronner.com direct-to-consumer, Amazon, retailer e-commerce platforms)")
add_bullet(doc, "International markets (15+ countries including Europe, Japan, Australia, and beyond)")
add_bullet(doc, "Institutional and bulk buyers (cleaning, hospitality)")

doc.add_heading("Geographic Footprint", level=2)
doc.add_paragraph(
    "Dr. Bronner's geographic footprint spans the entire United States domestically, with particular density "
    "in natural and organic retail channels. Key geographic considerations include:"
)
add_bullet(doc, "Manufacturing and HQ: ", "Vista, California (San Diego County) \u2014 ")
add_bullet(doc, "International operations across 14+ markets and 40+ countries via the All-One International Initiative")
add_bullet(doc, "Fair trade supply chain projects in Ghana, Sri Lanka, Samoa, India, Kenya, Palestine/Israel, Colombia, and Brazil")
add_bullet(doc, "European operations including a new Dr. Bronner's Haus and Museum in Laupheim, Germany")
add_bullet(doc, "Distribution to nearly 10,000 retail stores nationwide for core products")

doc.add_heading("Freight Profile Summary", level=2)
doc.add_paragraph(
    "Dr. Bronner's freight profile is characterized by a high volume of liquid and bar soap products moving "
    "from a single manufacturing origin in Vista, CA to distribution points and retail customers nationwide "
    "and internationally. The product mix creates a freight-dense profile with liquid goods predominating, "
    "resulting in significant weight per pallet. The company likely ships a mix of full truckloads (FTL) to "
    "major retail distribution centers (Walmart, Target, Costco, Kroger) and less-than-truckload (LTL) "
    "shipments to smaller natural retailers, specialty stores, and regional distributors."
)

doc.add_heading("Priority 1 Value Opportunities", level=2)
add_bullet(doc, "LTL consolidation and rate optimization for shipments to smaller natural and specialty retailers")
add_bullet(doc, "FTL capacity and pricing for seasonal surges and large retail DC shipments")
add_bullet(doc, "Cabo TMS platform for real-time visibility across a complex, multi-channel distribution network")
add_bullet(doc, "Carrier vetting and sustainability alignment (SmartWay partnership)")
add_bullet(doc, "Expedited shipping capability for time-sensitive promotional or seasonal launches")
add_bullet(doc, "Cross-border support for Canada/Mexico shipments")
add_bullet(doc, "Freight bill audit and cost analysis to identify savings opportunities")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 2: PRODUCTS & FREIGHT PROFILE TRANSLATION
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("2. Products and Freight Profile Translation", level=1)
add_thin_line(doc)

doc.add_heading("Product Characteristics Affecting Shipping", level=2)
doc.add_paragraph(
    "Dr. Bronner's product portfolio has several characteristics that directly impact freight planning, "
    "carrier selection, and cost management:"
)
add_bullet(doc, "Liquid-Heavy Product Mix: ", "The majority of Dr. Bronner's revenue comes from liquid castile soap in various sizes (8 oz to 1 gallon). Liquids are heavy relative to cube, meaning shipments will frequently weigh out before cubing out. This affects freight class and carrier pricing.")
add_bullet(doc, "Glass-Free Packaging: ", "Products are packaged in 100% post-consumer recycled (PCR) plastic bottles and FSC-certified refill cartons. No glass means lower breakage claims risk, but pallets of liquid still require careful stacking and securing to prevent shifting.")
add_bullet(doc, "Concentrated Formulas: ", "Dr. Bronner's products are concentrated, meaning more product per bottle and higher weight density per case. This is favorable from a shipping efficiency standpoint but increases weight per pallet.")
add_bullet(doc, "Temperature Sensitivity: ", "While soap products are generally not temperature-controlled, extreme cold can affect viscosity of liquid soaps and extreme heat can soften bar soaps. Seasonal routing considerations may apply for certain lanes.")
add_bullet(doc, "Multiple SKU Complexity: ", "With 8+ scent varieties across multiple product categories and size formats, order profiles can be complex with mixed-SKU pallets, increasing warehouse pick-and-pack complexity and potential for LTL shipments with varied item types.")

doc.add_heading("LTL vs. FTL Mix Analysis", level=2)
doc.add_paragraph(
    "Based on Dr. Bronner's distribution strategy and customer base, the estimated freight mode split is:"
)

make_table(doc, ["Freight Mode", "Estimated Share", "Primary Use Case"], [
    ["Full Truckload (FTL)", "50-60%", "Large replenishment orders to major retail DCs (Walmart, Target, Costco, Kroger, Whole Foods regional DCs)"],
    ["Less-Than-Truckload (LTL)", "30-40%", "Smaller orders to natural grocers, specialty retailers, regional distributors, and independent stores"],
    ["Parcel / Small Package", "5-10%", "Direct-to-consumer e-commerce orders via drbronner.com"],
    ["International Ocean/Air", "Variable", "Finished goods to 15+ international markets; raw material imports from fair trade suppliers"],
])

doc.add_paragraph()
doc.add_paragraph(
    "The LTL segment represents a significant opportunity for Priority1, as Dr. Bronner's services thousands "
    "of smaller retail accounts across the natural and specialty channel. These shipments are likely frequent, "
    "multi-stop, and require reliable transit times to meet retailer compliance requirements."
)

doc.add_heading("Palletization and Density Considerations", level=2)
doc.add_paragraph(
    "Liquid soap cases are dense and heavy. A standard 48x40 pallet of 32 oz liquid castile soap cases "
    "could easily weigh 1,800-2,200 lbs depending on case count and stacking height. Key considerations:"
)
add_bullet(doc, "Pallets will frequently exceed density thresholds, making weight-based pricing more favorable than dimensional")
add_bullet(doc, "High-density shipments may qualify for lower freight classes if properly classified")
add_bullet(doc, "Mixed pallets (liquid + bar soap + toothpaste) will have varying densities, requiring careful class determination")
add_bullet(doc, "Gallon-size liquid soap shipments (institutional/bulk) will be especially heavy per pallet")

doc.add_heading("NMFC / Freight Class Sensitivity", level=2)
doc.add_paragraph(
    "Soap and personal care products generally fall within NMFC classes 60-85, depending on packaging, "
    "density, and specific product type. Key classification considerations:"
)
add_bullet(doc, "Liquid castile soap in plastic bottles: Likely NMFC Class 60-70 (high density, low damage risk)")
add_bullet(doc, "Bar soap cases: Likely NMFC Class 65-77.5 (moderate density, compact packaging)")
add_bullet(doc, "Toothpaste and lip balm: May classify higher due to smaller packaging and lower density per case")
add_bullet(doc, "Proper classification and density-based re-classification can yield significant savings")
add_bullet(doc, "Priority1's freight bill audit capability can identify misclassification and recover overcharges")

doc.add_heading("Accessorial Exposure", level=2)
doc.add_paragraph(
    "Dr. Bronner's shipping profile likely exposes the company to several common accessorials:"
)
add_bullet(doc, "Liftgate delivery: ", "For smaller natural retail and specialty store deliveries lacking dock access")
add_bullet(doc, "Inside delivery: ", "Some specialty retailers may require inside delivery for small orders")
add_bullet(doc, "Limited access delivery: ", "Deliveries to downtown urban natural food stores or co-ops")
add_bullet(doc, "Residential delivery: ", "If any D2C fulfillment uses LTL for larger orders")
add_bullet(doc, "Appointment / notification: ", "Major retail DCs (Walmart, Target) require strict delivery appointments with fines for non-compliance")
add_bullet(doc, "Reweigh and reclassification: ", "Given the density of liquid products, carrier reweighs are common and can result in unexpected charges")

doc.add_heading("Where Priority1 Adds Value", level=2)
add_bullet(doc, "Rate optimization through group purchasing power across 47,000+ carrier partners")
add_bullet(doc, "NMFC classification expertise to ensure proper freight class and avoid overpayment")
add_bullet(doc, "Accessorial management and pre-audit to prevent surprise charges")
add_bullet(doc, "Cabo TMS for consolidated shipment management across FTL and LTL modes")
add_bullet(doc, "Carrier selection algorithms that balance cost, transit time, and service quality")
add_bullet(doc, "Freight bill audit and dispute resolution for post-shipment cost recovery")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 3: SUPPLY CHAIN & NETWORK MAPPING
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("3. Supply Chain and Network Mapping", level=1)
add_thin_line(doc)

doc.add_heading("Upstream Suppliers and Raw Material Sourcing", level=2)
doc.add_paragraph(
    "Dr. Bronner's supply chain is uniquely complex due to its commitment to organic, fair trade, and "
    "regenerative organic certified ingredients sourced from around the world. The company operates through "
    "its sister LLC, Serendiworld, which manages vertically integrated fair trade supply projects:"
)

make_table(doc, ["Ingredient", "Source Region", "Supply Project", "Scale"], [
    ["Coconut Oil", "Sri Lanka", "Serendipol", "1,200+ farmers, 20,000 acres, 30M coconuts/year"],
    ["Coconut Oil", "Samoa", "SerendiCoco Samoa", "Smallholder farmer network"],
    ["Palm Oil", "Ghana", "Serendipalm", "Fair trade certified palm oil"],
    ["Palm Kernel Oil", "Colombia", "Partner project", "Organic and fair trade certified"],
    ["Olive Oil", "Palestine / Israel", "Partner projects", "Fair trade certified"],
    ["Mint Oil", "India", "Serendimenthe", "Organic peppermint and spearmint"],
    ["Essential Oils", "Kenya", "SerendiKenya", "Fair trade essential oils"],
    ["Cane Sugar", "Brazil", "Partner project", "Organic fair trade certified"],
])

doc.add_paragraph()
doc.add_paragraph(
    "A total of 18,122 smallholder farmers contribute to Dr. Bronner's fair trade and organic supply chains "
    "globally. Over 75,000 acres of agricultural supply chain land have transitioned to Regenerative Organic "
    "Certified (ROC) status. To date, these projects have invested over $2.5 million in fair trade premiums, "
    "benefiting more than 20,000 people directly and indirectly."
)

doc.add_heading("Manufacturing", level=2)
doc.add_paragraph(
    "All manufacturing takes place at Dr. Bronner's facility in Vista, California. The company moved to this "
    "location in 2013 from a smaller facility in Escondido, CA, gaining roughly 8-10x more space. The Vista "
    "plant features:"
)
add_bullet(doc, "Bulk liquid castile soap storage tanks ('the tank farm') color-coded by scent")
add_bullet(doc, "Automated production lines for liquid soap filling, labeling, and packaging")
add_bullet(doc, "Bar soap manufacturing and packaging lines")
add_bullet(doc, "Toothpaste, balm, and personal care product manufacturing")
add_bullet(doc, "On-site warehouse operations for finished goods staging")
doc.add_paragraph(
    "The single-origin manufacturing model means all outbound domestic freight originates from the Vista, CA "
    "area, creating a hub-and-spoke distribution pattern with significant volume moving eastbound across the country."
)

doc.add_heading("Warehousing and Distribution Strategy", level=2)
doc.add_paragraph(
    "Dr. Bronner's operates warehouse operations at or near its Vista, CA headquarters. Based on its distribution "
    "footprint serving nearly 10,000 retail stores nationwide plus international markets, the company likely "
    "utilizes a combination of:"
)
add_bullet(doc, "On-site finished goods warehouse at the Vista facility")
add_bullet(doc, "Possible third-party distribution centers in strategic locations (East Coast, Midwest) to reduce transit times and costs for cross-country shipments")
add_bullet(doc, "Direct-to-retail-DC shipping for major accounts (Walmart, Target, Costco, Kroger)")
add_bullet(doc, "Distributor partnerships for natural channel retailers (UNFI, KeHE are common natural products distributors)")
add_bullet(doc, "Product donation warehouse for community outreach programs")

doc.add_heading("Downstream Customers and Channels", level=2)
doc.add_paragraph(
    "Dr. Bronner's finished goods flow through multiple distribution channels, each with distinct freight requirements:"
)

make_table(doc, ["Channel", "Key Customers", "Freight Characteristics"], [
    ["Mass Retail", "Walmart, Target, Costco, Kroger", "FTL to regional DCs; strict compliance requirements; appointment-based delivery"],
    ["Natural Grocery", "Whole Foods, Sprouts, Natural Grocers, NCG co-ops", "Mix of FTL (regional DCs) and LTL (store-direct); frequent replenishment cycles"],
    ["Specialty Retail", "REI, ULTA Beauty, Trader Joe's", "LTL and FTL depending on volume; may have unique packaging/labeling requirements"],
    ["Drug / Pharmacy", "Walgreens", "FTL to DCs; compliance-driven"],
    ["Distributors", "UNFI, KeHE (likely)", "FTL to distributor warehouses; distributor handles last-mile to stores"],
    ["E-Commerce (D2C)", "drbronner.com, Amazon", "Parcel-dominant; possible LTL for bulk replenishment to fulfillment centers"],
    ["International", "15+ markets, 40+ countries", "Ocean freight for finished goods; coordination with international logistics providers"],
])

doc.add_heading("Sustainability in Transportation", level=2)
doc.add_paragraph(
    "Sustainability is central to Dr. Bronner's brand identity and operations. The company has made significant "
    "commitments that directly impact logistics and transportation decisions:"
)
add_bullet(doc, "Zero-waste commitment: ", "Dr. Bronner's set a goal of becoming a zero-waste company, reducing packaging materials and shipping waste")
add_bullet(doc, "Concentrated formulas: ", "Products are formulated as concentrates, reducing the volume and weight of shipments per unit of product")
add_bullet(doc, "100% recycled packaging: ", "PCR plastic bottles and recycled cardboard shipping materials reduce environmental footprint")
add_bullet(doc, "Refill cartons: ", "New 32 oz refill cartons use 80% less plastic than traditional bottles, reducing packaging weight in transit")
add_bullet(doc, "Carbon footprint awareness: ", "As a company deeply committed to environmental stewardship, Dr. Bronner's likely evaluates carrier partners on emissions performance")

doc.add_paragraph(
    "Priority1's SmartWay partnership and ability to provide carriers with strong environmental credentials "
    "is a significant alignment point. Priority1 can offer Dr. Bronner's visibility into the carbon footprint "
    "of their freight operations and help identify lower-emission carrier options."
)

doc.add_heading("Transportation Risk Points", level=2)
doc.add_paragraph(
    "Key points where Priority1 can reduce cost or complexity in Dr. Bronner's supply chain:"
)
add_bullet(doc, "Single-origin risk: ", "All manufacturing from Vista, CA means any disruption (natural disaster, capacity crunch in Southern California) affects 100% of outbound freight. Priority1's 47,000+ carrier network provides resilience and alternative options.")
add_bullet(doc, "Long-haul exposure: ", "Significant freight volume must travel cross-country from Southern California to East Coast and Midwest markets. Priority1 can optimize mode selection (intermodal vs. OTR) for cost savings on these lanes.")
add_bullet(doc, "Seasonal demand spikes: ", "Holiday gifting seasons and promotional launches at major retailers create capacity needs. Priority1's carrier relationships ensure capacity availability during peak periods.")
add_bullet(doc, "Retail compliance: ", "Major retailers impose strict delivery windows, MABD (Must Arrive By Date) requirements, and chargeback penalties. Priority1's tracking and appointment management reduce compliance risk.")
add_bullet(doc, "International inbound complexity: ", "Raw materials arriving from 8+ countries require coordination with ocean freight and drayage. Priority1's drayage capabilities can support port-to-plant movements.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 4: INDUSTRY-SPECIFIC LOGISTICS CHALLENGES
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("4. Industry-Specific Logistics Challenges", level=1)
add_thin_line(doc)

doc.add_paragraph(
    "The consumer packaged goods (CPG) and natural personal care industry presents specific logistics "
    "challenges that affect companies like Dr. Bronner's. Below we identify these challenges and explain "
    "how Priority1's services directly address them."
)

doc.add_heading("Capacity Volatility", level=2)
doc.add_paragraph(
    "The CPG sector is subject to significant capacity fluctuations driven by seasonal demand, promotional "
    "cycles, and broader economic conditions. Dr. Bronner's, shipping primarily from Southern California, "
    "faces additional capacity pressure from port congestion in the LA/Long Beach corridor and competition "
    "for outbound trucking capacity with the massive volume of consumer goods flowing from West Coast ports."
)
doc.add_paragraph(
    "Priority1 Solution: With 47,000+ vetted carrier partners across all modes, Priority1 provides reliable "
    "capacity even during tight market conditions. The Cabo TMS enables instant quoting across multiple "
    "carriers, ensuring Dr. Bronner's always has options. Priority1's volume-based relationships with national "
    "and regional LTL carriers guarantee space allocation and competitive rates year-round."
)

doc.add_heading("Rate Pressure and Cost Management", level=2)
doc.add_paragraph(
    "Freight costs represent a significant line item for any CPG company shipping heavy liquid products "
    "across the country. Rate volatility in both truckload spot markets and LTL tariff increases can erode "
    "margins, particularly for a company like Dr. Bronner's that prioritizes paying fair wages and investing "
    "in sustainable supply chains."
)
doc.add_paragraph(
    "Priority1 Solution: Priority1's group purchasing power, aggregating 1.5+ million shipments annually, "
    "secures discounted rates that individual shippers cannot achieve independently. For Dr. Bronner's, this "
    "means access to enterprise-level pricing on LTL shipments to thousands of retail locations. Priority1's "
    "freight bill audit service also identifies overcharges and recovers costs post-shipment."
)

doc.add_heading("Service Failures and On-Time Delivery", level=2)
doc.add_paragraph(
    "In the retail CPG space, on-time delivery is critical. Major retailers like Walmart and Target impose "
    "chargebacks for late or non-compliant deliveries. Even in the natural channel, retailers like Whole Foods "
    "and Sprouts have strict receiving windows. Service failures can result in out-of-stocks, lost sales, "
    "damaged retailer relationships, and financial penalties."
)
doc.add_paragraph(
    "Priority1 Solution: Priority1's Cabo TMS provides real-time tracking and proactive exception management. "
    "Dedicated account representatives monitor shipments and can intervene early when delays are detected. "
    "Priority1's expedited shipping division offers backup options when standard transit is at risk, ensuring "
    "Dr. Bronner's meets MABD requirements consistently."
)

doc.add_heading("Accessorial Exposure", level=2)
doc.add_paragraph(
    "CPG companies shipping to diverse retail environments face significant accessorial charges. Deliveries "
    "to natural food co-ops in urban areas, small independent retailers without docks, and specialty stores "
    "in limited-access locations all trigger additional fees that can add 15-25% to base freight costs if "
    "not managed proactively."
)
doc.add_paragraph(
    "Priority1 Solution: Priority1 proactively identifies accessorial-prone deliveries during the quoting "
    "process, building these costs into upfront pricing rather than allowing surprise charges. The Cabo TMS "
    "maintains delivery location profiles that flag liftgate, limited access, and appointment requirements "
    "automatically, reducing billing surprises and improving cost predictability."
)

doc.add_heading("Claims Risk", level=2)
doc.add_paragraph(
    "While Dr. Bronner's products are packaged in durable plastic (no glass), liquid products are still "
    "susceptible to leakage from cap failure, case crushing from improper stacking, and label damage from "
    "moisture or friction during transit. Claims processing is time-consuming and can strain carrier relationships."
)
doc.add_paragraph(
    "Priority1 Solution: Priority1's carrier vetting process prioritizes service quality and claims ratios. "
    "The company's operations team assists with claims filing and resolution, reducing the administrative "
    "burden on Dr. Bronner's logistics staff. Priority1 can also recommend packaging and palletization best "
    "practices based on experience with similar CPG shippers."
)

doc.add_heading("Visibility and Data", level=2)
doc.add_paragraph(
    "For a company shipping to 10,000+ retail locations across multiple channels, visibility into shipment "
    "status, carrier performance, and cost analytics is essential. Many mid-sized CPG companies lack "
    "enterprise TMS platforms and rely on manual tracking, spreadsheets, or fragmented carrier portals."
)
doc.add_paragraph(
    "Priority1 Solution: The Cabo TMS provides a single platform for quoting, booking, tracking, and "
    "analyzing all shipments across modes. API/EDI integrations can connect with Dr. Bronner's existing "
    "systems (ERP, WMS) for seamless data flow. Advanced reporting and analytics help Dr. Bronner's identify "
    "cost trends, carrier performance issues, and optimization opportunities across their network."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 5: RECENT NEWS & TRIGGER EVENTS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("5. Recent News and Trigger Events Analysis", level=1)
add_thin_line(doc)

doc.add_paragraph(
    "The following recent developments at Dr. Bronner's may create logistics complexity and opportunities "
    "for Priority1 to provide value."
)

doc.add_heading("Living Wage Certification (January 2026)", level=3)
doc.add_paragraph(
    "Dr. Bronner's became a Certified Living Wage employer in January 2026, with a starting hourly wage of "
    "$29.11 (72% above California minimum wage). This reflects the company's commitment to worker welfare "
    "but also increases operating costs, potentially making freight cost optimization more critical to "
    "maintaining margins."
)
doc.add_paragraph(
    "Opportunity: Position Priority1 as a partner that helps offset rising labor costs through freight "
    "savings and operational efficiency."
)

doc.add_heading("Dropped B Corp Certification (February 2025)", level=3)
doc.add_paragraph(
    "Dr. Bronner's announced it would not renew its B Corp certification, citing concerns about the "
    "integrity of the certification when shared with large multinationals. The company co-launched the "
    "Purpose Pledge framework as an alternative accountability mechanism."
)
doc.add_paragraph(
    "Opportunity: Dr. Bronner's is actively seeking authentic sustainability partners. Priority1's SmartWay "
    "partnership and genuine commitment to responsible logistics aligns with Dr. Bronner's values-driven "
    "approach to vendor selection."
)

doc.add_heading("2025 Label Refresh and Refill Carton National Launch", level=3)
doc.add_paragraph(
    "Dr. Bronner's redesigned labels for its core soap products and launched soap refill cartons into "
    "nearly 10,000 retail stores nationwide. This represents a significant distribution expansion that "
    "increases outbound shipping volume and complexity."
)
doc.add_paragraph(
    "Opportunity: A national product launch of this scale requires reliable, scalable freight solutions. "
    "Priority1 can support the increased LTL volume to thousands of retail locations while maintaining "
    "cost discipline."
)

doc.add_heading("Manufacturing Facility Expansion", level=3)
doc.add_paragraph(
    "Dr. Bronner's has undergone significant expansion of its Vista, CA manufacturing facility to support "
    "growing production volumes. Increased production capacity translates directly to increased outbound "
    "freight volume."
)
doc.add_paragraph(
    "Opportunity: Growing production means growing freight spend. Priority1 can help Dr. Bronner's scale "
    "its transportation operations efficiently as production capacity expands."
)

doc.add_heading("International Market Expansion", level=3)
doc.add_paragraph(
    "The All-One International Initiative, now in its seventh year, operates across 14 markets and donated "
    "$238,710 outside the U.S. in 2024 (over $1 million total to date). The opening of the Dr. Bronner's "
    "Haus and Museum in Laupheim, Germany signals growing European investment."
)
doc.add_paragraph(
    "Opportunity: International growth increases the complexity of inbound raw material logistics and "
    "outbound finished goods distribution. Priority1's cross-border capabilities and future international "
    "freight offerings can support this expansion."
)

doc.add_heading("Exit from Chocolate Business (2025)", level=3)
doc.add_paragraph(
    "Dr. Bronner's exited the chocolate business and established a new parent company structure. This "
    "corporate restructuring may create opportunities as the logistics function refocuses on core soap "
    "and personal care products."
)
doc.add_paragraph(
    "Opportunity: Corporate restructuring often triggers a review of existing vendor relationships, "
    "including logistics providers. This is an ideal time for Priority1 to enter the conversation."
)

doc.add_heading("Regenerative Organic Certification Growth", level=3)
doc.add_paragraph(
    "Over 75,000 acres of Dr. Bronner's agricultural supply chain have transitioned to Regenerative "
    "Organic Certified status, and the company donated $8.4 million in charitable contributions in 2024."
)
doc.add_paragraph(
    "Opportunity: Emphasize Priority1's alignment with sustainability values and ability to provide "
    "environmentally responsible freight solutions."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 6: COMPETITIVE LANDSCAPE & BROKER OPPORTUNITY
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("6. Competitive Landscape and Broker Opportunity", level=1)
add_thin_line(doc)

doc.add_heading("Likely Current Logistics Provider Landscape", level=2)
doc.add_paragraph(
    "Based on Dr. Bronner's size (~$228M revenue), product type, and distribution profile, the company "
    "likely uses a combination of the following logistics approaches:"
)

make_table(doc, ["Provider Type", "Likely Usage", "Common Providers"], [
    ["Direct LTL Carrier Contracts", "High", "FedEx Freight, XPO, Estes, Old Dominion, SAIA for high-volume lanes"],
    ["Full Truckload Broker", "Moderate", "C.H. Robinson, Echo, TQL, or regional brokers for spot and overflow capacity"],
    ["Natural Products Distributor", "High", "UNFI, KeHE for distribution to natural channel retailers"],
    ["Parcel Carrier", "Moderate", "UPS, FedEx, USPS for e-commerce D2C fulfillment"],
    ["International Freight Forwarder", "Moderate", "For raw material imports and finished goods exports to 40+ countries"],
    ["Drayage Provider", "Moderate", "For port-to-plant movements of imported raw materials through LA/Long Beach"],
])

doc.add_heading("Common Incumbent Shortcomings", level=2)
doc.add_paragraph(
    "Companies of Dr. Bronner's size and profile often encounter the following pain points with existing "
    "logistics providers:"
)
add_bullet(doc, "Fragmented provider relationships: ", "Using multiple carriers and brokers without a unified platform creates visibility gaps and administrative overhead.")
add_bullet(doc, "Lack of proactive service: ", "Large 3PLs (C.H. Robinson, Echo) may treat a $228M company as a mid-tier account, providing reactive rather than proactive service.")
add_bullet(doc, "Rate complacency: ", "Long-standing carrier contracts may not reflect current market conditions. Without regular benchmarking, Dr. Bronner's may be overpaying on established lanes.")
add_bullet(doc, "Limited technology integration: ", "Many brokers offer basic tracking but lack the comprehensive TMS experience that Cabo provides.")
add_bullet(doc, "Sustainability blind spots: ", "Most logistics providers pay lip service to sustainability but cannot provide meaningful data on carrier emissions or environmental performance.")
add_bullet(doc, "Accessorial surprise charges: ", "Without proactive accessorial management, invoices routinely exceed quoted prices.")

doc.add_heading("Priority1 Competitive Positioning", level=2)
doc.add_paragraph(
    "Priority1 is uniquely positioned to outperform likely incumbent providers for Dr. Bronner's business. "
    "The following positioning statements can be used in sales conversations:"
)

doc.add_heading("vs. Large 3PLs (C.H. Robinson, Echo, TQL)", level=3)
doc.add_paragraph(
    '"At Priority1, Dr. Bronner\'s will never be just another account number. Our model pairs dedicated '
    'account representatives with proprietary technology to deliver the personal attention of a boutique '
    'provider backed by the scale and carrier access of a billion-dollar logistics company. We manage 1.5 '
    'million shipments annually \u2014 we have the buying power \u2014 but we deploy it with a personal touch that '
    'larger 3PLs simply cannot match."'
)

doc.add_heading("vs. Direct Carrier Contracts", level=3)
doc.add_paragraph(
    '"Direct carrier contracts may seem cost-effective, but they limit flexibility and create single points '
    'of failure. Priority1 gives Dr. Bronner\'s access to 47,000+ vetted carriers through a single platform, '
    'ensuring you always have capacity and competitive rates. Our group purchasing power means we can often '
    'beat direct contract rates, especially on LTL, while providing the visibility and management tools that '
    'individual carrier portals cannot."'
)

doc.add_heading("vs. Regional Brokers", level=3)
doc.add_paragraph(
    '"Regional brokers may know their local lanes, but Dr. Bronner\'s ships nationwide from Vista, CA to '
    '10,000+ retail locations across the country. Priority1 combines national carrier coverage with local '
    'expertise through our 50+ offices and independent agent network. We offer the geographic breadth to '
    'handle every lane in your network with consistent service and pricing."'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 7: SALES STRATEGY & MESSAGING FRAMEWORK
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("7. Sales Strategy and Messaging Framework", level=1)
add_thin_line(doc)

doc.add_heading("Target Contacts", level=2)
make_table(doc, ["Title / Role", "Relevance", "Approach"], [
    ["Director of Logistics", "Primary decision-maker for carrier selection and freight management", "Lead with operational value: technology, visibility, cost savings, carrier performance"],
    ["VP of Operations", "Oversees manufacturing and warehouse operations; influences logistics strategy", "Lead with efficiency gains, capacity reliability, and alignment with production scaling"],
    ["VP of Sales", "Manages retail customer relationships; cares about on-time delivery and compliance", "Lead with service reliability, retail DC compliance expertise, and fewer chargebacks"],
    ["VP of Finance", "Controls budgets and evaluates vendor ROI; has managed facility expansion financials", "Lead with cost savings, freight spend analytics, and total cost of ownership reduction"],
    ["Sustainability / CSR Lead", "Evaluates vendors on environmental and social criteria", "Lead with SmartWay partnership, carrier sustainability credentials, and emissions reporting"],
])

doc.add_heading("Likely Pain Points", level=2)
add_bullet(doc, "Rising freight costs eating into margins, especially as the company invests heavily in fair wages and sustainability programs")
add_bullet(doc, "Capacity challenges shipping from Southern California, particularly during peak season and port congestion periods")
add_bullet(doc, "Managing delivery compliance for major retail accounts (Walmart, Target, Costco) with strict MABD and appointment requirements")
add_bullet(doc, "Lack of unified visibility across multiple carriers and shipment modes")
add_bullet(doc, "Administrative burden of managing freight for 10,000+ retail delivery points across multiple channels")
add_bullet(doc, "Finding logistics partners whose values align with Dr. Bronner's commitment to sustainability and ethical business")
add_bullet(doc, "Scaling logistics operations to match growing production capacity and international expansion")

doc.add_heading("Discovery Questions", level=2)
doc.add_paragraph("The following questions are designed to uncover Dr. Bronner's freight needs and pain points:")
add_bullet(doc, '"How do you currently manage your LTL shipments to natural and specialty retail accounts? Are you working with a single provider or multiple carriers?"')
add_bullet(doc, '"What does your freight mix look like between FTL and LTL? Has that changed as your retail distribution has expanded?"')
add_bullet(doc, '"How are you handling capacity during peak seasons? Have you experienced challenges securing trucks out of Southern California?"')
add_bullet(doc, '"What technology are you using to manage freight? Do you have a TMS, or are you working with carrier portals individually?"')
add_bullet(doc, '"How important is sustainability in your logistics vendor selection? Do you evaluate carriers on emissions performance?"')
add_bullet(doc, '"What are your biggest headaches with retail DC compliance? Have chargebacks been an issue?"')
add_bullet(doc, '"With your recent product launches and manufacturing expansion, how has your outbound freight volume changed?"')
add_bullet(doc, '"Are you satisfied with the level of visibility you have into shipment status across your network?"')
add_bullet(doc, '"How do you currently handle freight cost analysis and benchmarking? Do you feel confident you are getting the best rates?"')

doc.add_heading("Priority1 Value Propositions for Dr. Bronner's", level=2)

doc.add_heading("1. Cost Savings Through Scale", level=3)
doc.add_paragraph(
    "Priority1 manages 1.5+ million shipments annually, giving us purchasing power that translates directly "
    "to lower rates for our customers. For Dr. Bronner's LTL shipments to thousands of retail locations, "
    "our group rates consistently outperform direct carrier contracts and smaller brokers. We estimate "
    "potential LTL savings of 10-20% based on typical CPG shipper profiles."
)

doc.add_heading("2. Technology That Simplifies", level=3)
doc.add_paragraph(
    "Our proprietary Cabo TMS provides a single platform to quote, book, track, and analyze all shipments "
    "across LTL and FTL modes. For a company like Dr. Bronner's with complex, multi-channel distribution, "
    "Cabo eliminates the need to log into multiple carrier portals and provides consolidated reporting "
    "and analytics. API/EDI integrations connect seamlessly with existing ERP and WMS systems."
)

doc.add_heading("3. Reliability and Capacity", level=3)
doc.add_paragraph(
    "With 47,000+ vetted carrier partners, Priority1 ensures Dr. Bronner's always has capacity, even during "
    "peak seasons and market disruptions. Our dedicated account team proactively monitors shipments and "
    "intervenes before problems escalate, protecting Dr. Bronner's retail relationships and avoiding "
    "costly chargebacks."
)

doc.add_heading("4. Values Alignment", level=3)
doc.add_paragraph(
    "Priority1 is a SmartWay partner committed to responsible logistics. We understand that Dr. Bronner's "
    "selects partners who share its commitment to sustainability and ethical business. Our carrier vetting "
    "process, emissions-conscious routing options, and transparent business practices align with the values "
    "that make Dr. Bronner's a category leader."
)

doc.add_heading("5. Personalized Service Model", level=3)
doc.add_paragraph(
    "Unlike large 3PLs where mid-sized accounts receive generic service, Priority1's model ensures Dr. "
    "Bronner's receives dedicated attention from experienced logistics professionals who understand the "
    "CPG space. Our 'Your product, our priority' philosophy means proactive communication, fast issue "
    "resolution, and a true partnership approach."
)

doc.add_heading("Proof Points", level=2)
add_bullet(doc, "$1.3 billion in annual revenue demonstrating financial stability and market credibility")
add_bullet(doc, "1.5+ million shipments managed annually with access to 47,000+ carrier partners")
add_bullet(doc, "80,000+ customers served across diverse industries including food, beverage, and consumer goods")
add_bullet(doc, "Proprietary Cabo TMS with instant quoting, real-time tracking, and freight bill audit capabilities")
add_bullet(doc, "50+ offices across the U.S. and Canada providing local expertise with national reach")
add_bullet(doc, "SmartWay partnership for sustainability-conscious freight management")
add_bullet(doc, "Named among Top 50 logistics providers in North America by Transport Topics")

doc.add_heading("Objection Handling", level=2)

make_table(doc, ["Objection", "Response"], [
    ['"We already have carrier contracts in place."',
     '"We respect your existing relationships. Many of our customers started by letting us handle overflow or specific lanes where we could demonstrate value. We can run a no-obligation rate comparison on your top lanes to show the savings potential without disrupting your current operations."'],
    ['"We are happy with our current broker."',
     '"That is great to hear. We find that even satisfied shippers benefit from a second set of eyes on their freight spend. Our freight audit and benchmarking analysis is complimentary and often uncovers 10-20% in savings that current providers have not identified. Can we run a quick analysis for you?"'],
    ['"We do not have time to onboard a new provider."',
     '"We understand. Priority1 is designed for easy onboarding. Our Cabo TMS can be set up in days, not weeks, and our team handles the heavy lifting. Many customers start with just a few lanes and expand as they see results."'],
    ['"We need a provider who understands sustainability."',
     '"Absolutely. Priority1 is a SmartWay partner, and we take sustainability seriously. We can provide carrier-level emissions data, help optimize routes for lower carbon impact, and ensure your logistics partners meet the same standards your brand represents."'],
    ['"Your company is not big enough for our needs."',
     '"Priority1 manages over $1.3 billion in freight annually and serves 80,000+ customers. We have the scale and carrier access to handle your volume with ease. But unlike the largest 3PLs, we combine that scale with personalized attention \u2014 your account will never be an afterthought."'],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 8: EXECUTIVE-LEVEL LOGISTICS SUMMARY
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("8. Executive-Level Logistics Summary", level=1)
add_thin_line(doc)

doc.add_heading("Company at a Glance", level=2)

make_table(doc, ["Attribute", "Detail"], [
    ["Company", "Dr. Bronner's Magic Soaps"],
    ["Headquarters", "Vista, California"],
    ["Revenue (est.)", "~$228 million"],
    ["Employees", "300-500"],
    ["CEO", "David Bronner"],
    ["Industry", "Natural Personal Care / CPG"],
    ["Products", "Organic castile soaps, toothpaste, lip balm, lotions, household cleaner"],
    ["Manufacturing", "Single facility in Vista, CA"],
    ["Distribution", "~10,000 retail locations; 40+ countries; major retailers + natural channel"],
    ["Key Differentiator", "Organic, fair trade, regenerative certified; values-driven brand"],
])

doc.add_heading("Strategic Shipping Risks", level=2)
add_bullet(doc, "Single manufacturing origin: ", "100% of production from Vista, CA creates concentration risk for outbound freight and makes the company highly dependent on West Coast carrier capacity and rates.")
add_bullet(doc, "Long-haul cost exposure: ", "Shipping heavy liquid products from Southern California to East Coast and Midwest markets results in significant per-unit transportation costs that directly impact product margins.")
add_bullet(doc, "Retail compliance risk: ", "Expanding mass-market distribution (Walmart, Target, Costco) brings increasingly stringent delivery requirements with financial penalties for non-compliance.")
add_bullet(doc, "Growth-driven complexity: ", "Rapid expansion of product lines, retail partnerships, and international markets is increasing logistics complexity faster than a small logistics team can manage efficiently.")
add_bullet(doc, "Supply chain sustainability pressure: ", "As a brand built on ethical sourcing and environmental responsibility, Dr. Bronner's faces scrutiny on transportation emissions and must ensure logistics partners meet high standards.")

doc.add_heading("Key Cost Drivers", level=2)
add_bullet(doc, "Outbound LTL freight to thousands of smaller retail accounts (natural grocery, specialty, co-ops)")
add_bullet(doc, "Full truckload shipments to major retail distribution centers across the country")
add_bullet(doc, "Cross-country fuel surcharges on long-haul lanes from Southern California")
add_bullet(doc, "Accessorial charges for deliveries to diverse retail environments (liftgate, limited access, appointments)")
add_bullet(doc, "Inbound raw material logistics from international fair trade suppliers via ocean and drayage")
add_bullet(doc, "E-commerce fulfillment and parcel shipping for direct-to-consumer orders")

doc.add_heading("Growth Impact on Logistics", level=2)
doc.add_paragraph(
    "Dr. Bronner's has grown from $4M to $228M in revenue over 25 years, with significant acceleration "
    "in the last decade. The company's recent manufacturing expansion, national refill carton launch into "
    "10,000 stores, label refresh, and international market development signal continued growth that will "
    "increase freight volume and complexity. Key growth drivers affecting logistics include:"
)
add_bullet(doc, "Expansion into mass retail channels (Walmart, Target, Costco) requiring higher-volume, compliance-driven FTL shipments")
add_bullet(doc, "National launch of refill cartons adding new SKUs and distribution points to the network")
add_bullet(doc, "International expansion into 14+ markets requiring finished goods export logistics")
add_bullet(doc, "Growing e-commerce/D2C business increasing small-parcel and LTL fulfillment volume")
add_bullet(doc, "Continued expansion of personal care product lines beyond soap (toothpaste, lotion, hair care)")

doc.add_heading("Why Priority1 as a Strategic Logistics Partner", level=2)
doc.add_paragraph(
    "Dr. Bronner's needs a logistics partner that can scale with its growth, deliver cost savings without "
    "compromising service, and align with its values-driven business philosophy. Priority1 is that partner "
    "because:"
)
add_bullet(doc, "We deliver enterprise-level pricing and carrier access through $1.3B in managed freight spend, while maintaining the personalized, high-touch service that a values-driven company like Dr. Bronner's deserves.")
add_bullet(doc, "Our proprietary Cabo TMS provides the technology platform to manage complex, multi-channel distribution from a single interface \u2014 critical as Dr. Bronner's logistics operation scales.")
add_bullet(doc, "Priority1 is a privately held, growth-oriented company that shares Dr. Bronner's entrepreneurial spirit and long-term perspective, unlike PE-backed or publicly traded 3PLs focused on quarterly results.")
add_bullet(doc, "Our SmartWay partnership, carrier vetting standards, and commitment to responsible logistics align with Dr. Bronner's brand values and sustainability commitments.")
add_bullet(doc, "With 50+ offices, 47,000+ carriers, and dedicated account management, we provide the reliability and resilience that Dr. Bronner's needs as it grows from a natural channel leader to a mass-market CPG powerhouse.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# APPENDIX: INTERNAL DEAL BRIEF (ONE PAGE)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Internal Deal Brief: Dr. Bronner's Magic Soaps", level=1)
add_thin_line(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("CONFIDENTIAL \u2014 Priority1 Internal Use Only  |  March 2026")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Arial'

doc.add_heading("Company Snapshot", level=2)
make_table(doc, ["Field", "Detail"], [
    ["Company Name", "Dr. Bronner's Magic Soaps"],
    ["Headquarters", "Vista, CA (San Diego County)"],
    ["Revenue", "~$228M (estimated)"],
    ["Employees", "300-500"],
    ["Industry", "Natural Personal Care / CPG"],
    ["CEO", "David Bronner"],
    ["Key Products", "Castile liquid soap, bar soap, toothpaste, lip balm, body lotion, Sal Suds cleaner"],
    ["Retail Presence", "~10,000 stores (Walmart, Target, Costco, Whole Foods, Sprouts, Kroger, REI, Walgreens)"],
    ["Manufacturing", "Single facility, Vista, CA"],
    ["International", "40+ countries, 14+ active markets"],
])

doc.add_heading("Freight Profile", level=2)
make_table(doc, ["Dimension", "Assessment"], [
    ["Primary Origin", "Vista, CA (all outbound)"],
    ["Est. Annual Freight Spend", "$8M-$15M+ (estimated 4-7% of revenue for heavy CPG)"],
    ["Mode Split", "50-60% FTL / 30-40% LTL / 5-10% parcel"],
    ["Product Density", "High (liquid soap is heavy; pallets weigh out before cubing out)"],
    ["Freight Class", "NMFC 60-85 (soap, personal care)"],
    ["Key Lanes", "Vista, CA to nationwide retail DCs and distributors; heavy East/Midwest lanes"],
    ["Seasonality", "Holiday gifting peaks; promotional launch surges"],
    ["Accessorial Risk", "Moderate-High (liftgate, limited access, retail appointments)"],
])

doc.add_heading("Risks and Opportunities", level=2)

make_table(doc, ["Risks", "Opportunities"], [
    ["Single-origin manufacturing limits routing flexibility", "LTL consolidation savings for natural/specialty retail shipments"],
    ["Heavy liquid products drive high freight costs per unit", "Freight class optimization and density-based reclassification"],
    ["Strict retail compliance requirements (Walmart, Target)", "Cabo TMS for visibility, tracking, and appointment management"],
    ["Values-driven vendor selection may require extra qualification", "SmartWay partnership and sustainability alignment differentiation"],
    ["May have strong incumbent carrier relationships", "Freight audit and benchmarking to demonstrate savings potential"],
])

doc.add_heading("Recommended Service Focus", level=2)
add_bullet(doc, "LTL (Primary): ", "High-volume LTL to natural grocery, specialty, and co-op retailers nationwide. This is Priority1's core strength and the segment most likely underserved by current providers.")
add_bullet(doc, "FTL (Secondary): ", "Truckload capacity for large retail DC shipments and seasonal surges. Competitive spot and contract rates from Vista, CA outbound.")
add_bullet(doc, "Freight Management: ", "Cabo TMS deployment for consolidated visibility and analytics across all modes. Freight bill audit to recover overcharges.")
add_bullet(doc, "Expedited (As Needed): ", "Time-critical shipments for promotional launches and MABD compliance recovery.")

doc.add_heading("Primary Sales Angle", level=2)
doc.add_paragraph(
    "Lead with values alignment and LTL expertise. Dr. Bronner's is a mission-driven company that evaluates "
    "partners on more than price. Position Priority1 as the logistics partner that combines meaningful cost "
    "savings with the personal service, technology, and sustainability credentials that Dr. Bronner's requires. "
    "Open the conversation with a complimentary freight benchmarking analysis on their top LTL lanes to "
    "demonstrate tangible savings potential without requiring commitment."
)

doc.add_heading("Next Steps", level=2)
make_table(doc, ["Step", "Action", "Timeline"], [
    ["1", "Identify and reach out to Director of Logistics and/or VP of Operations", "Week 1"],
    ["2", "Offer complimentary freight spend analysis and LTL rate benchmarking", "Week 1-2"],
    ["3", "Schedule introductory call; present Priority1 capabilities and values alignment", "Week 2-3"],
    ["4", "Obtain sample shipment data or BOLs for rate comparison", "Week 3-4"],
    ["5", "Deliver savings analysis and Cabo TMS demo", "Week 4-5"],
    ["6", "Propose pilot program on select LTL lanes (natural channel shipments)", "Week 5-6"],
    ["7", "Execute pilot; measure and report results vs. incumbent providers", "Week 6-10"],
    ["8", "Expand scope based on pilot performance", "Week 10+"],
])

# ── Save Document ─────────────────────────────────────────────────────
output_path = "/home/user/bronnerdeepdive/Dr_Bronners_Customer_Analysis_Priority1.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
